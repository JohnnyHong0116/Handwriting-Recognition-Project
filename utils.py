import os
from docx import Document

def save_to_word_doc(text_content, filename="output.docx"):
    full_path = os.path.join(os.getcwd(), filename)

    # Check if the document already exists
    if os.path.exists(full_path):
        # Open the existing document
        doc = Document(full_path)
    else:
        # Create a new document
        doc = Document()
    
    # Append the new content
    doc.add_paragraph(text_content)
    
    # Save the document
    doc.save(full_path)
    print(f"Text saved to {full_path}")
